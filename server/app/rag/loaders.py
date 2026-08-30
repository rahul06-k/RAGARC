import os
import re
from typing import List, Dict, Any, Tuple


def clean_text(text: str) -> str:
    """Clean extracted text while preserving paragraphs and punctuation."""
    if not text:
        return ""
    # Normalize unicode whitespace
    text = re.sub(r'[\r\f\v]', '\n', text)
    # Remove null bytes and non-printable characters (except normal newlines/tabs)
    text = "".join(ch for ch in text if ch.isprintable() or ch in ('\n', '\t'))
    # Replace repeated tabs or spaces with a single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Replace 3 or more newlines with 2 newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


class DocumentPage:
    def __init__(self, page_number: int, text: str, metadata: Dict[str, Any] = None):
        self.page_number = page_number
        self.text = text
        self.metadata = metadata or {}


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> Tuple[List[DocumentPage], int]:
        """
        Load a file and return a list of DocumentPage objects with page numbers and total page count.
        Supports PDF, DOCX, and TXT files.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif ext == ".docx":
            return DocumentLoader._load_docx(file_path)
        elif ext == ".txt":
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: PDF, DOCX, TXT")

    @staticmethod
    def _load_pdf(file_path: str) -> Tuple[List[DocumentPage], int]:
        pages = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            total_pages = len(doc)
            for i, page in enumerate(doc):
                text = clean_text(page.get_text())
                if text:
                    pages.append(DocumentPage(page_number=i + 1, text=text, metadata={"total_pages": total_pages}))
            doc.close()
            return pages, total_pages
        except Exception as e:
            # Fallback to pypdf if PyMuPDF fails or is not available
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text = clean_text(page.extract_text() or "")
                    if text:
                        pages.append(DocumentPage(page_number=i + 1, text=text, metadata={"total_pages": total_pages}))
                return pages, total_pages
            except Exception as inner_e:
                raise RuntimeError(f"Failed to extract PDF text: {str(e)} / {str(inner_e)}")

    @staticmethod
    def _load_docx(file_path: str) -> Tuple[List[DocumentPage], int]:
        pages = []
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                cleaned = clean_text(para.text)
                if cleaned:
                    full_text.append(cleaned)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
                    if row_text:
                        full_text.append(row_text)

            combined = "\n\n".join(full_text)
            # Estimate pages (~2500 characters per page)
            chars_per_page = 2500
            if not combined:
                return [], 1
            
            chunks = [combined[i:i + chars_per_page] for i in range(0, len(combined), chars_per_page)]
            total_pages = max(1, len(chunks))
            for i, chunk in enumerate(chunks):
                pages.append(DocumentPage(page_number=i + 1, text=clean_text(chunk), metadata={"total_pages": total_pages}))
            return pages, total_pages
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX text: {str(e)}")

    @staticmethod
    def _load_txt(file_path: str) -> Tuple[List[DocumentPage], int]:
        try:
            content = ""
            for encoding in ("utf-8", "latin-1", "cp1252"):
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            content = clean_text(content)
            if not content:
                return [], 1

            chars_per_page = 2500
            chunks = [content[i:i + chars_per_page] for i in range(0, len(content), chars_per_page)]
            total_pages = max(1, len(chunks))
            pages = [DocumentPage(page_number=i + 1, text=clean_text(chunk), metadata={"total_pages": total_pages}) for i, chunk in enumerate(chunks)]
            return pages, total_pages
        except Exception as e:
            raise RuntimeError(f"Failed to extract TXT text: {str(e)}")
